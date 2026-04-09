import os
import sys

# Force UTF-8 on Windows for robust file/font handling
os.environ["PYTHONUTF8"] = "1"
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
import shutil
import random
from typing import List, Dict, Set

# Add project root to path to import backend
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

import datasets
from datasets import load_dataset
# Use HF Mirror for better connectivity in Vietnam
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
# Increase logging level for datasets to see why it hangs
datasets.utils.logging.set_verbosity_info()
from docx import Document
from fpdf import FPDF
from tqdm import tqdm
from backend.retrieval.vector_db import client as qdrant_client
from backend.config import settings

# --- CONFIGURATION ---
HF_REPO = "nhn309261/vietnamese-legal-documents"
OUTPUT_DIR = "legal_docs"
TARGET_TOTAL = 100
FORMAT_COUNTS = {
    "docx": 30,
    "pdf": 20,
    "txt": 50
}

# --- UTILS ---

def clean_and_setup_directories():
    """Wipes the legal_docs directory and recreates subfolders."""
    if os.path.exists(OUTPUT_DIR):
        print(f"🧹 Cleaning existing directory: {OUTPUT_DIR}")
        shutil.rmtree(OUTPUT_DIR)
    
    for fmt in FORMAT_COUNTS.keys():
        os.makedirs(os.path.join(OUTPUT_DIR, fmt), exist_ok=True)
    print(f"📁 Created subdirectories: {list(FORMAT_COUNTS.keys())}")

def get_documents_from_db() -> List[Dict]:
    """Fetches unique document metadata available in the local Qdrant collection."""
    print(f"📡 Connecting to Qdrant: {settings.QDRANT_COLLECTION}...")
    docs = {} # document_id -> metadata
    offset = None
    
    while True:
        hits, offset = qdrant_client.scroll(
            collection_name=settings.QDRANT_COLLECTION,
            limit=500,
            with_payload=True,
            with_vectors=False,
            offset=offset
        )
        for hit in hits:
            p = hit.payload
            d_id = p.get("document_id")
            if d_id and d_id not in docs:
                docs[d_id] = {
                    "id": d_id,
                    "number": p.get("document_number", "N/A"),
                    "type": p.get("legal_type", "Văn bản"),
                    "authority": p.get("issuing_authority", "Cơ quan ban hành"),
                    "sectors": p.get("legal_sectors", []),
                    "title": p.get("title", "Không tiêu đề")
                }
        
        if offset is None:
            break
            
    print(f"✅ Found {len(docs)} unique documents in database.")
    return list(docs.values())

def balanced_sampling(db_docs: List[Dict], target_total: int) -> List[Dict]:
    """
    Samples a target number of docs from the list, aiming for a balanced distribution.
    """
    if not db_docs:
        return []

    if len(db_docs) <= target_total:
        return db_docs
        
    # Cluster by type for balance
    groups = {}
    for d in db_docs:
        g = d.get('type', 'Văn bản')
        if g not in groups:
            groups[g] = []
        groups[g].append(d)
        
    selected = []
    group_names = list(groups.keys())
    
    while len(selected) < target_total and group_names:
        random.shuffle(group_names) # Shuffle group pick order
        for g_name in list(group_names):
            if groups[g_name]:
                selected.append(groups[g_name].pop(random.randrange(len(groups[g_name]))))
                if len(selected) >= target_total:
                    break
            else:
                group_names.remove(g_name)
    
    return selected

# --- FILE GENERATORS ---

def clean_filename(name: str) -> str:
    """Removes invalid characters for filenames."""
    # Replace slashes, colons, etc. commonly found in doc numbers
    return name.replace("/", "-").replace(":", "-").replace("*", "").replace("?", "").replace("\"", "").replace("<", "").replace(">", "").replace("|", "-").strip()

def get_filename(doc: Dict, ext: str) -> str:
    """Constructs filename: 'Nghị định 91-2015-NĐ-CP Chính phủ.txt'"""
    base = f"{doc['type']} {doc['number']} {doc['authority']}"
    return f"{clean_filename(base)}.{ext}"

def save_as_txt(doc: Dict, content: str):
    filename = get_filename(doc, "txt")
    path = os.path.join(OUTPUT_DIR, "txt", filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"TIÊU ĐỀ: {doc['title']}\n")
        f.write("="*20 + "\n\n")
        f.write(content)

def save_as_docx(doc: Dict, content: str):
    document = Document()
    document.add_heading(doc['title'], 0)
    document.add_paragraph(content)
    filename = get_filename(doc, "docx")
    path = os.path.join(OUTPUT_DIR, "docx", filename)
    document.save(path)

class PDF(FPDF):
    pass

def save_as_pdf(doc: Dict, content: str):
    """Generates a PDF using fpdf2 with Unicode support for Vietnamese."""
    # Ensure content is clean Unicode
    safe_content = content.encode('utf-8', errors='ignore').decode('utf-8')
    title = doc['title'].encode('utf-8', errors='ignore').decode('utf-8')
    
    pdf = FPDF()
    pdf.add_page()
    
    font_paths = [
        "C:/Windows/Fonts/Arial.ttf",
        "C:/Windows/Fonts/times.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    ]
    font_loaded = False
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                # fname= style= is the correct fpdf2 signature for TTF
                # Style can be empty for regular.
                pdf.add_font("CustomFont", style="", fname=fp)
                pdf.set_font("CustomFont", size=11)
                font_loaded = True
                break
            except Exception:
                continue
            
    if not font_loaded:
        pdf.set_font("Helvetica", size=11)
        
    pdf.multi_cell(0, 10, f"TIÊU ĐỀ: {title}\n\n")
    pdf.multi_cell(0, 8, safe_content)
    
    filename = get_filename(doc, "pdf")
    path = os.path.join(OUTPUT_DIR, "pdf", filename)
    
    try:
        pdf.output(path)
    except Exception as e:
        # Fallback if filename has problematic characters
        if "codec" in str(e):
             alt_filename = f"doc_{doc['id']}.pdf"
             path = os.path.join(OUTPUT_DIR, "pdf", alt_filename)
             pdf.output(path)
        else:
             raise e

# --- MAIN EXECUTION ---

def main():
    print("🚀 Starting Optimized Document Generation from DB & HF...")
    
    # 1. Setup folders
    clean_and_setup_directories()
    
    # 2. Get available documents from DB
    try:
        db_docs = get_documents_from_db()
        id_to_meta = {d['id']: d for d in db_docs}
    except Exception as e:
        print(f"❌ DB Connection Error: {e}")
        return

    if not db_docs:
        print("❌ No documents found in DB. Aborting.")
        return

    # 3. Opportunistically collect candidates from HF stream
    CANDIDATE_POOL_SIZE = 400 # Collect more than needed for better balance
    candidates = [] # List of (meta, content)
    
    print(f"📥 Collecting up to {CANDIDATE_POOL_SIZE} candidates from HF stream...")
    # Using streaming=True for fast start. If first call hangs, check HF connectivity.
    ds_content = load_dataset(HF_REPO, "content", split="data", streaming=True)
    
    records_processed = 0
    pbar_collect = tqdm(total=CANDIDATE_POOL_SIZE, desc="Collecting Candidates")
    
    try:
        for record in ds_content:
            records_processed += 1
            r_id = str(record['id']) # Ensure string comparison
            
            if r_id in id_to_meta:
                candidates.append({
                    "meta": id_to_meta[r_id],
                    "content": record['content']
                })
                pbar_collect.update(1)
                
            # Update bar info every 100 docs to show we're alive
            if records_processed % 50 == 0:
                pbar_collect.set_postfix({"scanned": records_processed, "matches": len(candidates)})
                
            if len(candidates) >= CANDIDATE_POOL_SIZE:
                break
    except Exception as e:
        print(f"\n⚠️ Error while streaming HF: {e}")
        print("💡 Tip: Check your internet connection or use HF_ENDPOINT=https://hf-mirror.com")
        
    pbar_collect.close()

    if not candidates:
        print(f"❌ No matching documents found in HF dataset after scanning {records_processed} records. Aborting.")
        print(f"   - DB IDs: {list(id_to_meta.keys())[:5]}... (total {len(id_to_meta)})")
        return

    print(f"✅ Collected {len(candidates)} candidates after scanning {records_processed} records.")

    # 4. Perform balanced sampling on the collected candidates
    print(f"⚖️ Performing balanced sampling for {TARGET_TOTAL} documents...")
    # Extract metadata only for sampling
    candidate_metas = [c['meta'] for c in candidates]
    selected_metas = balanced_sampling(candidate_metas, TARGET_TOTAL)
    
    # Create lookup for selected IDs to get their content
    selected_ids = {m['id'] for m in selected_metas}
    final_docs = [c for c in candidates if c['meta']['id'] in selected_ids]

    # 5. Assign formats to the final 100 documents
    random.shuffle(final_docs)
    
    cursor = 0
    generated_count = 0
    pbar_gen = tqdm(total=len(final_docs), desc="Generating Files")
    
    for fmt, count in FORMAT_COUNTS.items():
        for _ in range(count):
            if cursor >= len(final_docs):
                break
            
            doc_data = final_docs[cursor]
            meta = doc_data['meta']
            content = doc_data['content']
            
            try:
                if fmt == "txt":
                    save_as_txt(meta, content)
                elif fmt == "docx":
                    save_as_docx(meta, content)
                elif fmt == "pdf":
                    save_as_pdf(meta, content)
                generated_count += 1
            except Exception as e:
                print(f"⚠️ Error generating {fmt} for {meta['id']}: {e}")
            
            cursor += 1
            pbar_gen.update(1)
            
    pbar_gen.close()
    print(f"\n🎉 Successfully generated {generated_count} documents in '{OUTPUT_DIR}/'.")
    print(f"   - docx: {FORMAT_COUNTS['docx']} | pdf: {FORMAT_COUNTS['pdf']} | txt: {FORMAT_COUNTS['txt']}")

if __name__ == "__main__":
    main()
