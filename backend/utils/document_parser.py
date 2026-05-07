import os
import re
import subprocess
from typing import List, Dict, Any
import docx
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
try:
    import docx2txt
except ImportError:
    docx2txt = None

from backend.ingestion.chunker.core import chunker

class DocumentParser:
    def __init__(self):
        pass

    def extract_text_from_pdf(self, file_path: str) -> str:
        if not fitz:
            raise ImportError("PyMuPDF (fitz) is not installed. Run `pip install pymupdf`")
        text_content = []
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text_content.append(page.get_text())
            doc.close()
        except Exception as e:
            raise RuntimeError(f"Failed to read PDF {file_path}: {str(e)}")
        return "\n".join(text_content)

    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            text_content = [para.text for para in doc.paragraphs]
            return "\n".join(text_content)
        except Exception as e:
            raise RuntimeError(f"Failed to read DOCX {file_path}: {str(e)}")

    def extract_text_from_doc(self, file_path: str) -> str:
        """Extract text from legacy .doc (Word 97-2003) format.
        Tries: 1) docx2txt, 2) python-docx (some .doc files work), 3) raw binary extraction.
        """
        # Strategy 1: docx2txt (handles some .doc files)
        if docx2txt:
            try:
                text = docx2txt.process(file_path)
                if text and len(text.strip()) > 50:
                    return text
            except Exception:
                pass
        
        # Strategy 2: python-docx (works if the .doc is actually docx-compat)
        try:
            doc = docx.Document(file_path)
            text_content = [para.text for para in doc.paragraphs]
            text = "\n".join(text_content)
            if text and len(text.strip()) > 50:
                return text
        except Exception:
            pass

        # Strategy 3: Raw binary text extraction with strict filtering (last resort)
        try:
            with open(file_path, 'rb') as f:
                raw = f.read()
            
            # Extract only readable strings (4 or more printable chars)
            # This avoids most OLE binary structures
            import re
            
            # Try UTF-16-LE (common in Word)
            try:
                decoded = raw.decode('utf-16-le', errors='ignore')
                # Keep only printable Vietnamese & basic symbols
                # Filter: A-Z, a-z, 0-9, Vietnamese chars, and basic punctuation
                filtered = re.findall(r'[a-zA-Z0-9ÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚĂĐĨŨƠàáâãèéêìíòóôõùúăđĩũơƯĂÂÊÔƠƯàáâãèéêìíòóôõùúăđĩũơưăâêôơưẠ-ỹ\s\.\,\:\-\!\?]{10,}', decoded)
                text = "\n".join(filtered)
                if len(text.strip()) > 200:
                    return text
            except Exception:
                pass
                
            # Try UTF-8 fallback
            decoded = raw.decode('utf-8', errors='ignore')
            filtered = re.findall(r'[a-zA-Z0-9ÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚĂĐĨŨƠàáâãèéêìíòóôõùúăđĩũơƯĂÂÊÔƠƯàáâãèéêìíòóôõùúăđĩũơưăâêôơưẠ-ỹ\s\.\,\:\-\!\?]{10,}', decoded)
            text = "\n".join(filtered)
            if len(text.strip()) > 200:
                return text
        except Exception:
            pass

        raise RuntimeError(f"Failed to read .doc file {file_path}. Install docx2txt: pip install docx2txt")

    def parse_and_chunk(self, file_path: str, base_metadata: Dict[str, Any] = None) -> List[Dict]:
        """
        Extract text from file and apply Hierarchical Chunking (AdvancedLegalChunker).
        Ensures breadcrumbs like (Chương > Điều > Khoản) are preserved.
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".pdf":
            content = self.extract_text_from_pdf(file_path)
        elif ext == ".docx":
            content = self.extract_text_from_docx(file_path)
        elif ext == ".doc":
            content = self.extract_text_from_doc(file_path)
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        metadata = base_metadata or {}
        if "title" not in metadata:
            metadata["title"] = os.path.basename(file_path)

        # AdvancedLegalChunker handles Regex-based tree splitting and Breadcrumbs
        # skip_llm=True để xử lý nhanh cho file upload phiên chat, tránh gọi LLM bóc thực thể
        chunks = chunker.process_document(content, metadata, skip_llm=True)
        return chunks

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Trích xuất số hiệu văn bản từ 500 ký tự đầu tiên của file."""
        ext = os.path.splitext(file_path)[1].lower()
        content = ""
        try:
            if ext == ".pdf":
                content = self.extract_text_from_pdf(file_path)
            elif ext == ".docx":
                content = self.extract_text_from_docx(file_path)
            elif ext == ".doc":
                content = self.extract_text_from_doc(file_path)
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
        except Exception as e:
            print(f"      [Parser] Metadata extraction failed: {e}")

        # Lấy 500 ký tự đầu để tìm số hiệu chính
        preamble = content[:500]
        
        # 1. Tìm mẫu "Số: 123/QĐ-..." (Chính xác nhất cho tiêu đề văn bản)
        match_so = re.search(r"(?i)Số\s*:\s*([0-9]+/[0-9]{4}/[A-Z0-9Đ\-]+|[0-9]+/[A-Z0-9Đ\-]+)", preamble)
        if match_so:
            doc_number = match_so.group(1).strip()
        else:
            # 2. Fallback dùng hàm tổng quát
            from backend.ingestion.chunker import metadata as md
            doc_number = md.extract_doc_number(preamble) or "File Upload"

        return {
            "document_number": doc_number,
            "title": os.path.basename(file_path)
        }

parser = DocumentParser()
