import os
import sys
import io
import csv

# Force UTF-8 on Windows console to support emoji/Vietnamese
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

os.environ["PYTHONUTF8"] = "1"
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"

# HuggingFace cache → D: drive to avoid filling C:
os.environ["HF_HOME"] = r"D:\huggingface_cache"
os.environ["TRANSFORMERS_CACHE"] = r"D:\huggingface_cache"
os.environ["SENTENCE_TRANSFORMERS_HOME"] = r"D:\huggingface_cache\sentence_transformers"

import shutil
import random
from typing import List, Dict, Set

# Add project root to path to import backend
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

import datasets
from datasets import load_dataset
# Increase logging level for datasets to see why it hangs
datasets.utils.logging.set_verbosity_info()
from docx import Document
from fpdf import FPDF
from tqdm import tqdm

# --- CONFIGURATION ---
HF_REPO = "nhn309261/vietnamese-legal-docs"
CSV_PATH = os.path.join(PROJECT_ROOT, "metadata_the_thao_y_te_100.csv")
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "legal_docs")
OUTPUT_2026_DIR = os.path.join(OUTPUT_DIR, "2026")

TARGET_TOTAL = 100
FORMAT_COUNTS = {
    "pdf": 50,
    "docx": 50,
}

# 2026 docs config
TARGET_2026 = 5
SECTOR_FILTER = "Thể thao - Y tế"

# --- CSV READER ---

def load_metadata_from_csv(csv_path: str) -> List[Dict]:
    """
    Reads the metadata CSV file and returns a list of document metadata dicts.
    CSV columns: id, document_number, title, legal_type, promulgation_date, signer_name, legal_sectors
    """
    docs = []
    print(f"📄 Reading metadata from: {csv_path}")
    with open(csv_path, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            docs.append({
                "id": str(row["id"]).strip(),
                "number": row.get("document_number", "N/A").strip(),
                "type": row.get("legal_type", "Văn bản").strip(),
                "authority": row.get("signer_name", "").strip(),
                "sectors": row.get("legal_sectors", "").strip(),
                "title": row.get("title", "Không tiêu đề").strip(),
                "date": row.get("promulgation_date", "").strip(),
            })
    print(f"✅ Loaded {len(docs)} documents from CSV.")
    return docs


# --- UTILS ---

def clean_and_setup_directories():
    """Wipes the legal_docs directory and recreates subfolders for pdf/docx + 2026."""
    if os.path.exists(OUTPUT_DIR):
        print(f"🧹 Cleaning existing directory: {OUTPUT_DIR}")
        shutil.rmtree(OUTPUT_DIR)

    for fmt in FORMAT_COUNTS.keys():
        os.makedirs(os.path.join(OUTPUT_DIR, fmt), exist_ok=True)

    # Create 2026 subfolders
    for fmt in FORMAT_COUNTS.keys():
        os.makedirs(os.path.join(OUTPUT_2026_DIR, fmt), exist_ok=True)

    print(f"📁 Created subdirectories: {list(FORMAT_COUNTS.keys())} + 2026/")


# --- FILE GENERATORS ---

def clean_filename(name: str) -> str:
    """Removes invalid characters for filenames."""
    return (name.replace("/", "-").replace(":", "-").replace("*", "")
            .replace("?", "").replace("\"", "").replace("<", "")
            .replace(">", "").replace("|", "-").strip())


def get_filename(doc: Dict, ext: str) -> str:
    """Constructs filename: 'Nghị định 91-2015-NĐ-CP Cơ quan.ext'"""
    base = f"{doc['type']} {doc['number']} {doc['authority']}"
    # Truncate if too long for Windows path limits
    if len(base) > 120:
        base = base[:120]
    return f"{clean_filename(base)}.{ext}"


def save_as_docx(doc: Dict, content: str, output_base: str):
    """Generates a DOCX file with title + content."""
    document = Document()
    document.add_heading(doc['title'], 0)
    # Add metadata paragraph
    meta_text = (
        f"Số hiệu: {doc['number']}\n"
        f"Loại văn bản: {doc['type']}\n"
        f"Lĩnh vực: {doc['sectors']}\n"
        f"Ngày ban hành: {doc.get('date', 'N/A')}\n"
        f"Người ký: {doc['authority']}\n"
    )
    p = document.add_paragraph(meta_text)
    p.style = document.styles['Intense Quote']
    document.add_paragraph(content)
    filename = get_filename(doc, "docx")
    path = os.path.join(output_base, "docx", filename)
    document.save(path)
    return path


def save_as_pdf(doc: Dict, content: str, output_base: str):
    """Generates a PDF using fpdf2 with Unicode support for Vietnamese."""
    safe_content = content.encode('utf-8', errors='ignore').decode('utf-8')
    title = doc['title'].encode('utf-8', errors='ignore').decode('utf-8')

    pdf = FPDF()
    pdf.add_page()

    font_paths = [
        "C:/Windows/Fonts/Arial.ttf",
        "C:/Windows/Fonts/times.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    ]
    font_loaded = False
    for fp in font_paths:
        if os.path.exists(fp):
            try:
                pdf.add_font("CustomFont", style="", fname=fp)
                pdf.set_font("CustomFont", size=11)
                font_loaded = True
                break
            except Exception:
                continue

    if not font_loaded:
        pdf.set_font("Helvetica", size=11)

    # Title
    pdf.multi_cell(0, 10, f"TIÊU ĐỀ: {title}\n\n")
    # Metadata block
    meta_text = (
        f"Số hiệu: {doc['number']}\n"
        f"Loại văn bản: {doc['type']}\n"
        f"Lĩnh vực: {doc['sectors']}\n"
        f"Ngày ban hành: {doc.get('date', 'N/A')}\n"
        f"Người ký: {doc['authority']}\n"
    )
    pdf.multi_cell(0, 8, meta_text)
    pdf.multi_cell(0, 8, "=" * 40 + "\n")
    pdf.multi_cell(0, 8, safe_content)

    filename = get_filename(doc, "pdf")
    path = os.path.join(output_base, "pdf", filename)

    try:
        pdf.output(path)
    except Exception as e:
        if "codec" in str(e).lower():
            alt_filename = f"doc_{doc['id']}.pdf"
            path = os.path.join(output_base, "pdf", alt_filename)
            pdf.output(path)
        else:
            raise e
    return path


# --- MAIN: DOWNLOAD 100 DOCS FROM CSV ---

def download_csv_docs(csv_docs: List[Dict]):
    """
    Downloads content for 100 docs listed in the CSV from HuggingFace,
    saves as 50 PDF + 50 DOCX.
    """
    id_to_meta = {d['id']: d for d in csv_docs}
    target_ids = set(id_to_meta.keys())

    print(f"\n{'='*60}")
    print(f"📥 PHASE 1: Downloading {len(target_ids)} documents from CSV metadata")
    print(f"{'='*60}")
    print(f"   Source: {HF_REPO} (content config)")
    print(f"   Output: {OUTPUT_DIR}/pdf/ + {OUTPUT_DIR}/docx/")

    # Stream content from HuggingFace
    ds_content = load_dataset(HF_REPO, "content", split="data", streaming=True)

    matched = []  # List of (meta, content)
    scanned = 0
    pbar = tqdm(total=len(target_ids), desc="Matching CSV docs")

    try:
        for record in ds_content:
            scanned += 1
            r_id = str(record['id'])

            if r_id in target_ids:
                matched.append({
                    "meta": id_to_meta[r_id],
                    "content": record['content']
                })
                pbar.update(1)

            if scanned % 100 == 0:
                pbar.set_postfix({"scanned": scanned, "matched": len(matched)})

            # Stop early if we found all
            if len(matched) >= len(target_ids):
                break
    except Exception as e:
        print(f"\n⚠️ Error while streaming HF: {e}")
        print("💡 Tip: Check your internet connection or HF_ENDPOINT")

    pbar.close()
    print(f"✅ Matched {len(matched)}/{len(target_ids)} documents after scanning {scanned} records.\n")

    if not matched:
        print("❌ No matching documents found. Aborting Phase 1.")
        return

    # Shuffle and split into PDF / DOCX
    random.shuffle(matched)

    pdf_count = FORMAT_COUNTS["pdf"]
    docx_count = FORMAT_COUNTS["docx"]

    generated = {"pdf": 0, "docx": 0}
    errors = 0
    pbar_gen = tqdm(total=len(matched), desc="Generating PDF/DOCX files")

    for i, doc_data in enumerate(matched):
        meta = doc_data['meta']
        content = doc_data['content']

        # First half → PDF, second half → DOCX
        if generated["pdf"] < pdf_count:
            fmt = "pdf"
        else:
            fmt = "docx"

        try:
            if fmt == "pdf":
                save_as_pdf(meta, content, OUTPUT_DIR)
            else:
                save_as_docx(meta, content, OUTPUT_DIR)
            generated[fmt] += 1
        except Exception as e:
            errors += 1
            print(f"  ⚠️ Error generating {fmt} for {meta['id']} ({meta['number']}): {e}")

        pbar_gen.update(1)

    pbar_gen.close()
    total_gen = generated["pdf"] + generated["docx"]
    print(f"\n🎉 Phase 1 Complete: Generated {total_gen} files")
    print(f"   📄 PDF:  {generated['pdf']}")
    print(f"   📝 DOCX: {generated['docx']}")
    if errors > 0:
        print(f"   ⚠️  Errors: {errors}")


# --- MAIN: DOWNLOAD 5 DOCS FROM 2026 ---

def download_2026_docs():
    """
    Searches HuggingFace metadata for 2026 documents in the Thể thao - Y tế sector,
    then downloads their content and saves as PDF/DOCX in the 2026/ folder.
    """
    print(f"\n{'='*60}")
    print(f"📥 PHASE 2: Finding {TARGET_2026} documents from year 2026")
    print(f"{'='*60}")
    print(f"   Sector filter: {SECTOR_FILTER}")
    print(f"   Output: {OUTPUT_2026_DIR}/pdf/ + {OUTPUT_2026_DIR}/docx/\n")

    # Step 1: Stream metadata to find 2026 docs in the right sector
    print("🔍 Scanning HF metadata for 2026 documents...")
    ds_meta = load_dataset(HF_REPO, "metadata", split="data", streaming=True)

    candidates_2026 = []
    scanned = 0
    pbar = tqdm(desc="Scanning metadata for 2026 docs")

    try:
        for record in ds_meta:
            scanned += 1
            pbar.update(1)

            if scanned % 500 == 0:
                pbar.set_postfix({"scanned": scanned, "found": len(candidates_2026)})

            # Check year from issuance_date (format: DD/MM/YYYY)
            date_str = str(record.get("issuance_date", ""))
            sectors = str(record.get("legal_sectors", ""))

            # Parse year - handle both DD/MM/YYYY and YYYY-MM-DD formats
            year = None
            if "/" in date_str:
                parts = date_str.split("/")
                if len(parts) == 3:
                    year = parts[2]
            elif "-" in date_str:
                parts = date_str.split("-")
                if len(parts) == 3:
                    year = parts[0]

            if year == "2026" and SECTOR_FILTER in sectors:
                signers = record.get("signers", "")
                signer_name = ""
                if isinstance(signers, str) and ":" in signers:
                    signer_name = signers.split(":")[0]
                elif isinstance(signers, str):
                    signer_name = signers

                candidates_2026.append({
                    "id": str(record["id"]),
                    "number": record.get("document_number", "N/A"),
                    "type": record.get("legal_type", "Văn bản"),
                    "authority": signer_name,
                    "sectors": sectors,
                    "title": record.get("title", "Không tiêu đề"),
                    "date": date_str,
                })

            # Collect more than needed for variety, but stop eventually
            if len(candidates_2026) >= TARGET_2026 * 4:
                break
    except Exception as e:
        print(f"\n⚠️ Error scanning metadata: {e}")

    pbar.close()
    print(f"✅ Found {len(candidates_2026)} candidate 2026 documents after scanning {scanned} records.\n")

    if not candidates_2026:
        print("❌ No 2026 documents found in the specified sector. Skipping Phase 2.")
        return

    # Select up to TARGET_2026 documents
    if len(candidates_2026) > TARGET_2026:
        selected_2026 = random.sample(candidates_2026, TARGET_2026)
    else:
        selected_2026 = candidates_2026[:TARGET_2026]

    print(f"📋 Selected {len(selected_2026)} documents for download:")
    for i, doc in enumerate(selected_2026, 1):
        print(f"   {i}. [{doc['type']}] {doc['number']} - {doc['title'][:80]}...")

    # Step 2: Download content for selected 2026 docs
    target_ids = {d['id'] for d in selected_2026}
    id_to_meta = {d['id']: d for d in selected_2026}

    print(f"\n📥 Downloading content for {len(target_ids)} docs from HF content stream...")
    ds_content = load_dataset(HF_REPO, "content", split="data", streaming=True)

    matched_2026 = []
    scanned = 0
    pbar = tqdm(total=len(target_ids), desc="Downloading 2026 content")

    try:
        for record in ds_content:
            scanned += 1
            r_id = str(record['id'])

            if r_id in target_ids:
                matched_2026.append({
                    "meta": id_to_meta[r_id],
                    "content": record['content']
                })
                pbar.update(1)

            if scanned % 100 == 0:
                pbar.set_postfix({"scanned": scanned, "matched": len(matched_2026)})

            if len(matched_2026) >= len(target_ids):
                break
    except Exception as e:
        print(f"\n⚠️ Error downloading 2026 content: {e}")

    pbar.close()
    print(f"✅ Downloaded content for {len(matched_2026)}/{len(target_ids)} documents.\n")

    # Step 3: Save as PDF and DOCX (alternate)
    generated = {"pdf": 0, "docx": 0}
    for i, doc_data in enumerate(matched_2026):
        meta = doc_data['meta']
        content = doc_data['content']
        fmt = "pdf" if i % 2 == 0 else "docx"

        try:
            if fmt == "pdf":
                path = save_as_pdf(meta, content, OUTPUT_2026_DIR)
            else:
                path = save_as_docx(meta, content, OUTPUT_2026_DIR)
            generated[fmt] += 1
            print(f"   ✅ [{fmt.upper()}] {meta['number']} → saved")
        except Exception as e:
            print(f"   ⚠️ Error saving {fmt} for {meta['id']}: {e}")

    total_gen = generated["pdf"] + generated["docx"]
    print(f"\n🎉 Phase 2 Complete: Generated {total_gen} files in '{OUTPUT_2026_DIR}/'")
    print(f"   📄 PDF:  {generated['pdf']}")
    print(f"   📝 DOCX: {generated['docx']}")


# --- MAIN: EXPORT METADATA FOR SPLIT DATA ---

def export_split_metadata():
    """
    Exports metadata from HF repo as split CSV files for easy data management.
    Saves the full content+metadata join for the Thể thao - Y tế sector.
    """
    print(f"\n{'='*60}")
    print(f"📥 PHASE 3: Export split metadata from HF repo")
    print(f"{'='*60}")

    split_dir = os.path.join(PROJECT_ROOT, "datasets", "split_data")
    os.makedirs(split_dir, exist_ok=True)

    # Stream metadata and split by sector
    print("🔍 Streaming metadata from HF to split by sector...")
    ds_meta = load_dataset(HF_REPO, "metadata", split="data", streaming=True)

    all_meta = []
    the_thao_yte_meta = []
    scanned = 0
    pbar = tqdm(desc="Scanning metadata for split")

    try:
        for record in ds_meta:
            scanned += 1
            pbar.update(1)

            if scanned % 1000 == 0:
                pbar.set_postfix({"scanned": scanned, "thể thao - y tế": len(the_thao_yte_meta)})

            sectors = str(record.get("legal_sectors", ""))

            if SECTOR_FILTER in sectors:
                signers = record.get("signers", "")
                signer_name = ""
                if isinstance(signers, str) and ":" in signers:
                    signer_name = signers.split(":")[0]
                elif isinstance(signers, str):
                    signer_name = signers

                meta_row = {
                    "id": record["id"],
                    "document_number": record.get("document_number", ""),
                    "title": record.get("title", ""),
                    "legal_type": record.get("legal_type", ""),
                    "issuance_date": record.get("issuance_date", ""),
                    "signer_name": signer_name,
                    "legal_sectors": sectors,
                }
                the_thao_yte_meta.append(meta_row)

            # Limit total scan for efficiency
            if scanned >= 600000:
                break
    except Exception as e:
        print(f"\n⚠️ Error scanning metadata: {e}")

    pbar.close()

    # Save Thể thao - Y tế metadata
    output_csv = os.path.join(split_dir, "metadata_the_thao_y_te_full.csv")
    if the_thao_yte_meta:
        with open(output_csv, "w", encoding="utf-8", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=the_thao_yte_meta[0].keys())
            writer.writeheader()
            writer.writerows(the_thao_yte_meta)
        print(f"✅ Saved {len(the_thao_yte_meta)} Thể thao - Y tế records → {output_csv}")

    # Further split by year
    year_groups = {}
    for row in the_thao_yte_meta:
        date_str = str(row.get("issuance_date", ""))
        year = "unknown"
        if "/" in date_str:
            parts = date_str.split("/")
            if len(parts) == 3:
                year = parts[2]
        elif "-" in date_str:
            parts = date_str.split("-")
            if len(parts) == 3:
                year = parts[0]

        if year not in year_groups:
            year_groups[year] = []
        year_groups[year].append(row)

    for year, rows in sorted(year_groups.items()):
        year_csv = os.path.join(split_dir, f"metadata_the_thao_y_te_{year}.csv")
        with open(year_csv, "w", encoding="utf-8", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=rows[0].keys())
            writer.writeheader()
            writer.writerows(rows)
        print(f"   📊 Year {year}: {len(rows)} docs → {year_csv}")

    print(f"\n🎉 Phase 3 Complete: Split data saved to '{split_dir}/'")


# --- ENTRY POINT ---

def main():
    print("🚀 Legal Document Crawler - From CSV Metadata + HF Content")
    print(f"   HF Repo: {HF_REPO}")
    print(f"   CSV: {CSV_PATH}")
    print(f"   Output: {OUTPUT_DIR}\n")

    # Verify CSV exists
    if not os.path.exists(CSV_PATH):
        print(f"❌ CSV file not found: {CSV_PATH}")
        return

    # 1. Setup folders
    clean_and_setup_directories()

    # 2. Load metadata from CSV
    csv_docs = load_metadata_from_csv(CSV_PATH)
    if not csv_docs:
        print("❌ No documents found in CSV. Aborting.")
        return

    # 3. Phase 1: Download 100 docs from CSV → PDF/DOCX
    download_csv_docs(csv_docs)

    # 4. Phase 2: Download 5 docs from 2026 → 2026/pdf + 2026/docx
    download_2026_docs()

    # 5. Phase 3: Export split metadata from HF
    export_split_metadata()

    print(f"\n{'='*60}")
    print("✅ ALL PHASES COMPLETE!")
    print(f"{'='*60}")

    # Summary of output structure
    print(f"\n📂 Output Structure:")
    print(f"   {OUTPUT_DIR}/")
    print(f"   ├── pdf/          (50 files from CSV)")
    print(f"   ├── docx/         (50 files from CSV)")
    print(f"   └── 2026/")
    print(f"       ├── pdf/      (2-3 files from 2026)")
    print(f"       └── docx/     (2-3 files from 2026)")
    split_dir = os.path.join(PROJECT_ROOT, "datasets", "split_data")
    print(f"\n   {split_dir}/")
    print(f"       ├── metadata_the_thao_y_te_full.csv")
    print(f"       └── metadata_the_thao_y_te_<year>.csv (split by year)")


if __name__ == "__main__":
    main()
